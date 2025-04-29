import streamlit as st 
import requests 
from transformers import pipeline
from PyPDF2 import PdfReader 
import tempfile
import re
import requests 
import xml.etree.ElementTree as ET
import io
from docx import Document
from reportlab.pdfgen import canvas
import sqlite3
from transformers import pipeline




def search_arxiv(query):
    # Construct the URL with the search query
    url = f"http://export.arxiv.org/api/query?search_query={query}&start=0&max_results=3"
    
    # Send the request to the arXiv API
    response = requests.get(url)
    
    # Parse the XML response
    root = ET.fromstring(response.text)
    
    # Find all entries (papers)
    entries = root.findall('{http://www.w3.org/2005/Atom}entry')
    
    # List to store formatted paper details
    papers = []
    
    for entry in entries:
        # Extract paper details
        title = entry.find('{http://www.w3.org/2005/Atom}title').text
        authors = [author.find('{http://www.w3.org/2005/Atom}name').text for author in entry.findall('{http://www.w3.org/2005/Atom}author')]
        summary = entry.find('{http://www.w3.org/2005/Atom}summary').text
        published = entry.find('{http://www.w3.org/2005/Atom}published').text
        paper_link = entry.find('{http://www.w3.org/2005/Atom}link').attrib['href']
        pdf_link = entry.find('{http://www.w3.org/2005/Atom}link[@title="pdf"]').attrib['href']
        
        # Format and store the paper data
        paper_data = {
            'title': title,
            'authors': authors,
            'summary': summary,
            'published': published,
            'paper_link': paper_link,
            'pdf_link': pdf_link
        }
        papers.append(paper_data)
    
    return papers

# Code for download option
def create_docx(text):
    doc = Document()
    doc.add_heading("Summary", 0)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

def create_pdf(text):
    buffer = io.BytesIO()
    p = canvas.Canvas(buffer)
    text_object = p.beginText(40, 800)
    text_object.setFont("Helvetica", 12)
    for line in text.split('\n'):
        text_object.textLine(line)
    p.drawText(text_object)
    p.showPage()
    p.save()
    buffer.seek(0)
    return buffer

# Title
st.title("📚 AI-Powered Academic Research Assistant")

# Load summarizer and QA models
@st.cache_resource
def load_models():
    summarizer = pipeline("summarization")
    
    # Load both models
    extractive_qa = pipeline("question-answering", model="deepset/roberta-base-squad2", tokenizer="deepset/roberta-base-squad2")
    generative_qa = pipeline("text2text-generation", model="google/flan-t5-base")
    
    return summarizer, extractive_qa, generative_qa

# Database for login and history
def init_db():
    conn = sqlite3.connect("user_data.db")
    c = conn.cursor()
    # Create users table
    c.execute('''CREATE TABLE IF NOT EXISTS users (
        username TEXT PRIMARY KEY,
        password TEXT
    )''')

    # Create history table
    c.execute('''CREATE TABLE IF NOT EXISTS history (
        username TEXT,
        feature TEXT,
        input TEXT,
        output TEXT
    )''')
    
    conn.commit()
    conn.close()

init_db()

def save_history(username, feature, input_text, output_text):
    conn = sqlite3.connect("user_data.db")
    c = conn.cursor()
    c.execute("INSERT INTO history (username, feature, input, output) VALUES (?, ?, ?, ?)", 
              (username, feature, input_text, output_text))
    conn.commit()
    conn.close()


# PDF Text Extractor
def extract_text_from_pdf(pdf_file):
    reader = PdfReader(pdf_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() + "\n"
    return text

# Citation Generator (APA style)
def generate_apa_citation(title, authors, year, source):
    author_str = ", ".join(authors)
    return f"{author_str} ({year}). {title}. {source}."

def search_arxiv(query):
    # Construct the URL with the search query
    url = f"http://export.arxiv.org/api/query?search_query={query}&start=0&max_results=3"
    
    # Send the request to the arXiv API
    response = requests.get(url)
    
    # Parse the XML response
    root = ET.fromstring(response.text)
    
    # Find all entries (papers)
    entries = root.findall('{http://www.w3.org/2005/Atom}entry')
    
    # List to store formatted paper details
    papers = []
    
    for entry in entries:
        # Extract paper details
        title = entry.find('{http://www.w3.org/2005/Atom}title').text
        authors = [author.find('{http://www.w3.org/2005/Atom}name').text for author in entry.findall('{http://www.w3.org/2005/Atom}author')]
        summary = entry.find('{http://www.w3.org/2005/Atom}summary').text
        published = entry.find('{http://www.w3.org/2005/Atom}published').text
        paper_link = entry.find('{http://www.w3.org/2005/Atom}link').attrib['href']
        pdf_link = entry.find('{http://www.w3.org/2005/Atom}link[@title="pdf"]').attrib['href']
        
        # Format and store the paper data
        paper_data = {
            'title': title,
            'authors': authors,
            'summary': summary,
            'published': published,
            'paper_link': paper_link,
            'pdf_link': pdf_link
        }
        papers.append(paper_data)
    
    return papers

# Main options
st.sidebar.title("🔐 Login to Continue")

if "logged_in" not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.username = ""

# Login/Register Forms
if not st.session_state.logged_in:
    choice = st.sidebar.radio("Login or Register", ["Login", "Register"])
    username = st.sidebar.text_input("Username")
    password = st.sidebar.text_input("Password", type="password")

    if st.sidebar.button("Submit"):
        conn = sqlite3.connect("user_data.db")
        c = conn.cursor()

        if choice == "Register":
            try:
                c.execute("INSERT INTO users (username, password) VALUES (?, ?)", (username, password))
                conn.commit()
                st.sidebar.success("Registration successful. You can now login.")
            except sqlite3.IntegrityError:
                st.sidebar.error("Username already exists.")
        else:  # Login
            c.execute("SELECT * FROM users WHERE username=? AND password=?", (username, password))
            if c.fetchone():
                st.session_state.logged_in = True
                st.session_state.username = username
                st.sidebar.success(f"Welcome back, {username} 👋")
            else:
                st.sidebar.error("Invalid credentials")
        conn.close()
else:
    st.sidebar.success(f"Logged in as {st.session_state.username}")
    if st.sidebar.button("Logout"):
        st.session_state.logged_in = False
        st.session_state.username = ""

if st.session_state.logged_in:
    ("Search Papers")
    
    # Your existing logic for features (Search, Summarize, Ask) goes here...
    
else:
    st.warning("Please log in to use the features.")

#Main sidebar options after logging in
option = st.sidebar.selectbox("Choose Feature", ["Search Papers", "Summarize PDF", "Ask Questions", "📜 View History"])

if option == "Search Papers":
    query = st.text_input("Enter research topic")
    if st.button("Search"):
        with st.spinner("Searching arXiv..."):
            papers = search_arxiv(query)  # Call the function to get paper details
            
            # Check if papers were returned
            if papers:

                # Save search to history here:
                save_history(st.session_state.username, "Search Papers", query, f"{len(papers)} results found")

                # Loop through each paper and display its details
                for i, paper in enumerate(papers, 1):
                    # Display the title of the paper
                    st.subheader(f"{i}. {paper['title']}")
                    
                    # Display the authors' names
                    st.markdown(f"**Authors**: {', '.join(paper['authors'])}")
                    
                    # Display the date of publication
                    st.markdown(f"**Published on**: {paper['published']}")
                    
                    # Display the summary (abstract)
                    st.markdown(f"**Summary**: {paper['summary']}")
                    
                    # Provide links to the paper (abstract and PDF)
                    st.markdown(f"[Read Abstract]({paper['paper_link']}) | [Download PDF]({paper['pdf_link']})")
                    
                    # Add some space for clarity
                    st.write("\n")
            else:
                # If no papers were found, inform the user
                st.write("No results found.")

                if st.button("Search"):
                    save_history(st.session_state.username, "Search Papers", query, f"{len(papers)} results found")



elif option == "Summarize PDF":
    uploaded_file = st.file_uploader("Upload a research paper (PDF)", type="pdf")
    if uploaded_file is not None:
        reader = PdfReader(uploaded_file)
        total_pages = len(reader.pages)

        # Mode selection: Page Range vs Full Document
        mode = st.radio("Select summarization mode", ["By Page Range", "Entire PDF (Smart)"])

        if mode == "By Page Range":
            page_range = st.slider(
                f"Select page range (1 to {total_pages})",
                1, total_pages, (1, min(3, total_pages))
            )

            selected_text = ""
            for i in range(page_range[0] - 1, page_range[1]):
                selected_text += reader.pages[i].extract_text() + "\n"

            st.text_area("Extracted Text", selected_text[:3000], height=300)

            if st.button("Summarize Selected Pages"):
                try:
                    summary = summarizer(selected_text[:3000])[0]['summary_text']
                    st.subheader("Summary")
                    st.write(summary)
                    save_history(st.session_state.username, "Summarize PDF", selected_text[:3000], summary)
                    # Download buttons for selected pages summary
                    st.download_button("📄 Download Summary as TXT", data=summary, file_name="summary.txt")
                    st.download_button("📄 Download Summary as DOCX", data=create_docx(summary), file_name="summary.docx")
                    st.download_button("📄 Download Summary as PDF", data=create_pdf(summary), file_name="summary.pdf")

                except Exception as e:
                    st.error(f"Summarization failed: {e}")

        elif mode == "Entire PDF (Smart)":
            full_text = ""
            for page in reader.pages:
                full_text += page.extract_text() + "\n"

            st.text_area("Full PDF Extracted (Preview)", full_text[:2000], height=200)

            if st.button("Summarize Whole Document"):
                try:
                    summaries = []
                    for i in range(0, len(full_text), 3000):
                        chunk = full_text[i:i+3000]
                        summary = summarizer(chunk)[0]['summary_text']
                        summaries.append(summary)

                    final_summary = " ".join(summaries)
                    st.subheader("Smart Chunk Summary")
                    st.write(final_summary)
                    save_history(st.session_state.username, "Smart Summarize", full_text[:6000], final_summary)
                    # Download buttons for full document summary
                    st.download_button("📄 Download Smart Summary as TXT", data=final_summary, file_name="smart_summary.txt")
                    st.download_button("📄 Download Smart Summary as DOCX", data=create_docx(final_summary), file_name="smart_summary.docx")
                    st.download_button("📄 Download Smart Summary as PDF", data=create_pdf(final_summary), file_name="smart_summary.pdf")

                except Exception as e:
                    st.error(f"Smart summarization failed: {e}")

elif option == "Ask Questions":
    uploaded_file = st.file_uploader("Upload a research paper (PDF)", type="pdf")
    question = st.text_input("Ask a question from the paper")
    mode = st.radio("Answering Mode", ["Fast (Extractive)", "Smart (Generative)"])

    if uploaded_file is not None and question:
        reader = PdfReader(uploaded_file)
        full_text = ""
        for page in reader.pages:
            full_text += page.extract_text() + "\n"

        chunk_size = 2000  # Safer for both models
        chunks = [full_text[i:i+chunk_size] for i in range(0, len(full_text), chunk_size)]

        if st.button("Get Answer"):
            with st.spinner("Thinking... 🤔"):
                final_answer = "Sorry, no answer found."

                if mode == "Fast (Extractive)":
                    best_answer = {"answer": final_answer, "score": 0}
                    for chunk in chunks:
                        try:
                            result = extractive_qa(question=question, context=chunk)
                            if result["score"] > best_answer["score"]:
                                best_answer = result
                        except:
                            continue
                    final_answer = best_answer["answer"]

                elif mode == "Smart (Generative)":
                    all_answers = []
                    for chunk in chunks[:3]:  # Limit to 3 chunks for speed
                        prompt = f"Answer the question based on this text:\n{chunk}\n\nQuestion: {question}"
                        try:
                            output = generative_qa(prompt, max_length=100, do_sample=False)[0]["generated_text"]
                            all_answers.append(output.strip())
                        except:
                            continue
                    final_answer = max(all_answers, key=len) if all_answers else "No good answer generated."

            st.subheader("Answer")
            st.write(final_answer)
            save_history(st.session_state.username, "Ask Question", question, final_answer)

elif option == "📜 View History":
    st.header("📜 Your Activity History")
    conn = sqlite3.connect("user_data.db")
    c = conn.cursor()
    c.execute("SELECT feature, input, output FROM history WHERE username = ?", (st.session_state.username,))
    rows = c.fetchall()
    conn.close()

    if rows:
        for i, (feature, input_text, output_text) in enumerate(rows[::-1], 1):  # Show latest first
            with st.expander(f"{i}. {feature}"):
                st.markdown(f"**Input:**\n{input_text}")
                st.markdown(f"**Output:**\n{output_text}")
    else:
        st.info("No history yet. Use some features first!")
        